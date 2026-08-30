#!/usr/bin/env python3
"""Build the anonymous REF manuscript as a self-contained DOCX.

The source of record remains tex/theory.tex.  The conversion path is:

1. TeX4ht emits HTML with MathML.
2. MathML nodes are replaced by stable markers before LibreOffice conversion.
3. The converted Word document is normalized to the journal layout.
4. Markers are replaced by editable Office Math (OMML), and linked figures are
   embedded in the DOCX package.
"""

from __future__ import annotations

import argparse
import copy
import os
import re
import shutil
import subprocess
import tempfile
import urllib.parse
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree, html


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
R_DOC_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
XML_NS = "http://www.w3.org/XML/1998/namespace"
NS = {"w": W_NS, "m": M_NS}
MATH_MARKER = re.compile(r"\[\[MATH(\d{4})\]\]")
USABLE_WIDTH_DXA = 9360


def parse_args() -> argparse.Namespace:
    root = Path(__file__).resolve().parents[1]
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output",
        type=Path,
        default=root / "submission" / "ref_anonymous_manuscript.docx",
    )
    parser.add_argument("--soffice", type=Path)
    return parser.parse_args()


def run(command: list[str], *, cwd: Path, env: dict[str, str] | None = None) -> None:
    subprocess.run(command, cwd=cwd, env=env, check=True)


def locate_soffice(explicit: Path | None) -> Path:
    if explicit is not None:
        return explicit
    found = shutil.which("soffice")
    if found:
        return Path(found)
    bundled = Path(
        "/Users/akira/.cache/codex-runtimes/codex-primary-runtime/"
        "dependencies/bin/override/soffice"
    )
    if bundled.exists():
        return bundled
    raise FileNotFoundError("soffice was not found; pass --soffice explicitly")


def convert_tex_to_html(source_tex_dir: Path, work: Path) -> Path:
    copied_tex = work / "tex"
    shutil.copytree(source_tex_dir, copied_tex)
    html_dir = work / "html"
    html_dir.mkdir()
    run(
        ["make4ht", "-f", "html5", "-d", str(html_dir), "theory.tex", "mathml"],
        cwd=copied_tex,
    )
    output = html_dir / "theory.html"
    if not output.exists():
        raise RuntimeError("make4ht did not create theory.html")
    return output


def prepare_marked_html(source: Path) -> tuple[Path, list[bytes]]:
    document = html.parse(str(source))
    math_nodes = document.xpath('//*[local-name()="math"]')
    serialized: list[bytes] = []
    for index, node in enumerate(math_nodes):
        serialized.append(etree.tostring(node, encoding="utf-8", with_tail=False))
        marker = etree.Element("span")
        marker.text = f"[[MATH{index:04d}]]"
        marker.tail = node.tail
        node.getparent().replace(node, marker)

    for meta in document.xpath(
        '//meta[@name="generator" or @name="src" or @name="author"]'
    ):
        meta.getparent().remove(meta)

    heads = document.xpath("//head")
    if not heads:
        raise RuntimeError("TeX4ht HTML has no head element")
    style = etree.SubElement(heads[0], "style")
    style.text = """
@page { size: Letter portrait; margin: 1in; }
html, body {
  font-family: "Times New Roman", serif !important;
  font-size: 12pt !important;
  line-height: 1.5 !important;
}
h1 { font-size: 16pt !important; line-height: 1.15 !important; }
h2, h3, h4, h5 { font-size: 12pt !important; line-height: 1.2 !important; }
table.tabular {
  width: 100% !important;
  max-width: 100% !important;
  table-layout: fixed !important;
}
td, th {
  white-space: normal !important;
  overflow-wrap: anywhere !important;
}
img { max-width: 100% !important; height: auto !important; }
a { color: #000000 !important; text-decoration: none !important; }
"""
    marked = source.with_name("theory-marked.html")
    marked.write_bytes(
        etree.tostring(
            document,
            method="html",
            encoding="utf-8",
            doctype="<!DOCTYPE html>",
        )
    )
    return marked, serialized


def convert_html_to_docx(source: Path, work: Path, soffice: Path) -> Path:
    output_dir = work / "converted"
    output_dir.mkdir()
    profile = work / "lo-profile"
    profile.mkdir()
    lo_home = work / "lo-home"
    lo_home.mkdir()
    env = os.environ.copy()
    env["HOME"] = str(lo_home)
    env["TMPDIR"] = "/private/tmp"
    run(
        [
            str(soffice),
            "--headless",
            f"-env:UserInstallation=file://{profile}",
            "--convert-to",
            "docx:Office Open XML Text",
            "--outdir",
            str(output_dir),
            str(source),
        ],
        cwd=source.parent,
        env=env,
    )
    output = output_dir / f"{source.stem}.docx"
    if not output.exists():
        raise RuntimeError("LibreOffice did not create the intermediate DOCX")
    return output


def set_font(run, *, size: float = 12, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def style_kind(name: str) -> str:
    lowered = name.lower()
    if "titlehead" in lowered:
        return "title"
    if "abstracttitle" in lowered:
        return "abstract"
    if "sectionhead" in lowered and "subsectionhead" not in lowered:
        return "section"
    if "subsubsectionhead" in lowered:
        return "subsubsection"
    if "subsectionhead" in lowered:
        return "subsection"
    if "paragraphhead" in lowered:
        return "paragraph"
    if "bibliography" in lowered:
        return "bibliography"
    return "body"


def normalize_styles(document: Document) -> None:
    for style in document.styles:
        if not hasattr(style, "paragraph_format"):
            continue
        kind = style_kind(style.name)
        style.font.name = "Times New Roman"
        style._element.get_or_add_rPr().rFonts.set(
            qn("w:ascii"), "Times New Roman"
        )
        style._element.get_or_add_rPr().rFonts.set(
            qn("w:hAnsi"), "Times New Roman"
        )
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0, 0, 0)
        fmt = style.paragraph_format
        if kind == "title":
            style.font.bold = True
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(12)
            fmt.line_spacing = 1.5
        elif kind == "abstract":
            style.font.bold = True
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(6)
            fmt.line_spacing = 1.5
        elif kind == "section":
            style.font.bold = True
            fmt.space_before = Pt(12)
            fmt.space_after = Pt(6)
            fmt.line_spacing = 1.5
            fmt.keep_with_next = True
        elif kind == "subsection":
            style.font.bold = True
            fmt.space_before = Pt(10)
            fmt.space_after = Pt(4)
            fmt.line_spacing = 1.5
            fmt.keep_with_next = True
        elif kind in {"subsubsection", "paragraph"}:
            style.font.bold = True
            fmt.space_before = Pt(8)
            fmt.space_after = Pt(3)
            fmt.line_spacing = 1.5
            fmt.keep_with_next = True
        else:
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.line_spacing = 1.5
            if "indent" in style.name.lower():
                fmt.first_line_indent = Inches(0.25)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_margins(cell, *, top: int = 80, start: int = 80,
                     bottom: int = 80, end: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7B7B7")


def set_table_geometry(table, widths: list[int]) -> None:
    if len(widths) != len(table.columns) or sum(widths) != USABLE_WIDTH_DXA:
        raise ValueError("table widths must match the column count and sum to 9360")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def content_table_widths(header: str, columns: int) -> list[int]:
    normalized = " ".join(header.split())
    if normalized.startswith("Symbol"):
        return [1300, 2350, 2350, 3360]
    if normalized.startswith("Object"):
        return [1500, 2700, 2400, 2760]
    if normalized.startswith("Variable"):
        return [1400, 1900, 2000, 1800, 2260]
    if normalized.startswith("BOJ input / component"):
        return [1700, 1300, 2200, 2800, 1360]
    if normalized.startswith("Taxonomy identifier"):
        return [1800, 1500, 2400, 2400, 1260]
    if normalized.startswith("External source"):
        return [1500, 1400, 2100, 2100, 2260]
    if normalized.startswith("BOJ input / source"):
        return [1700, 1400, 2300, 2900, 1060]
    if normalized.startswith("BOJ series name / code") and columns == 5:
        return [2100, 1700, 1700, 1900, 1960]
    if normalized.startswith("BOJ series name / code") and columns == 6:
        return [1900, 1000, 1700, 1900, 1500, 1360]
    if normalized.startswith("BOJ series name / code") and columns == 10:
        return [1400, 1000, 900, 1050, 1050, 650, 950, 1050, 750, 560]
    if normalized.startswith("BOJ borrower group"):
        return [1700, 1100, 2100, 2900, 1560]
    if normalized.startswith("Audit"):
        return [1900, 2900, 4560]
    if columns == 8 and "Matched-stock RMSE" in normalized:
        return [400, 1450, 1550, 1300, 400, 900, 1350, 2010]
    if columns == 1:
        return [USABLE_WIDTH_DXA]
    if columns == 2:
        return [8360, 1000]
    base = USABLE_WIDTH_DXA // columns
    widths = [base] * columns
    widths[-1] += USABLE_WIDTH_DXA - sum(widths)
    return widths


def normalize_tables(document: Document) -> None:
    for table in document.tables:
        header = " | ".join(cell.text for cell in table.rows[0].cells)
        equation_table = len(table.rows) == 1 and "MATH" in header
        widths = content_table_widths(header, len(table.columns))
        set_table_geometry(table, widths)
        if not equation_table:
            set_table_borders(table)
            set_repeat_table_header(table.rows[0])
            for row in table.rows:
                prevent_row_split(row)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.5
                    for run in paragraph.runs:
                        set_font(run, size=12, bold=(row_index == 0 and not equation_table))


def split_long_audit_table(document: Document) -> None:
    """Split the two longest audit rows into a clean continued table."""
    for table in document.tables:
        if " ".join(table.rows[0].cells[0].text.split()) != "Audit":
            continue
        split_index = next(
            (
                index
                for index, row in enumerate(table.rows)
                if " ".join(row.cells[0].text.split()).startswith(
                    "Series availability"
                )
            ),
            None,
        )
        if split_index is None:
            return

        original = table._tbl
        continued = copy.deepcopy(original)
        original_rows = original.findall(qn("w:tr"))
        continued_rows = continued.findall(qn("w:tr"))
        for row in original_rows[split_index:]:
            original.remove(row)
        for row in continued_rows[1:split_index]:
            continued.remove(row)

        page_break = OxmlElement("w:p")
        properties = OxmlElement("w:pPr")
        style = OxmlElement("w:pStyle")
        style.set(qn("w:val"), "Normal")
        properties.append(style)
        properties.append(OxmlElement("w:pageBreakBefore"))
        page_break.append(properties)
        original.addnext(page_break)
        page_break.addnext(continued)
        return


def split_wide_source_table(document: Document) -> None:
    """Present the ten-column Appendix table as two editable continuations."""
    for table in list(document.tables):
        header = " ".join(table.rows[0].cells[0].text.split())
        if header != "BOJ series name / code" or len(table.columns) != 10:
            continue

        values = [[cell.text for cell in row.cells] for row in table.rows]
        left_columns = (0, 1, 2, 3, 4)
        right_columns = (0, 5, 6, 7, 8, 9)

        def make_table(indices: tuple[int, ...]):
            created = document.add_table(rows=len(values), cols=len(indices))
            for row_index, row_values in enumerate(values):
                for column_index, source_index in enumerate(indices):
                    created.cell(row_index, column_index).text = row_values[source_index]
            return created

        left = make_table(left_columns)
        continuation = document.add_paragraph("Table 8 (continued).")
        continuation.style = document.styles["Normal"]
        continuation.paragraph_format.keep_with_next = True
        continuation.paragraph_format.page_break_before = True
        right = make_table(right_columns)

        original = table._tbl
        original.addprevious(left._tbl)
        original.addprevious(continuation._p)
        original.addprevious(right._tbl)
        original.getparent().remove(original)
        return


def keep_captions_with_objects(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.replace("\xa0", " ").split())
        if text.startswith("Table "):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            if text.startswith(("Table 5:", "Table 6:")):
                paragraph.paragraph_format.page_break_before = True
            sibling = paragraph._p.getnext()
            while sibling is not None and sibling.tag == qn("w:p"):
                sibling_text = "".join(sibling.itertext()).strip()
                if sibling_text:
                    break
                properties = sibling.find(qn("w:pPr"))
                if properties is None:
                    properties = OxmlElement("w:pPr")
                    sibling.insert(0, properties)
                keep_next = properties.find(qn("w:keepNext"))
                if keep_next is None:
                    keep_next = OxmlElement("w:keepNext")
                    properties.append(keep_next)
                sibling = sibling.getnext()
        elif text.startswith("Figure "):
            paragraph.paragraph_format.keep_with_next = False
            paragraph.paragraph_format.keep_together = True
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.paragraph_format.keep_with_next = True


def resize_second_figure(document: Document) -> None:
    drawing_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph._p.xpath(".//w:drawing")
    ]
    if not drawing_paragraphs:
        return
    if len(drawing_paragraphs) != 2:
        raise RuntimeError(
            f"expected two figure paragraphs, found {len(drawing_paragraphs)}"
        )
    size_emu = str(round(4.35 * 914_400))
    second = drawing_paragraphs[1]._p
    for extent in second.xpath(".//wp:extent"):
        extent.set("cx", size_emu)
        extent.set("cy", size_emu)
    for extent in second.xpath(".//a:xfrm/a:ext"):
        extent.set("cx", size_emu)
        extent.set("cy", size_emu)


def format_hyperlink_runs(document: Document) -> None:
    for paragraph in document.paragraphs:
        for run in paragraph._p.xpath(".//w:hyperlink//w:r"):
            properties = run.find(qn("w:rPr"))
            if properties is None:
                properties = OxmlElement("w:rPr")
                run.insert(0, properties)
            run_style = properties.find(qn("w:rStyle"))
            if run_style is not None:
                properties.remove(run_style)
            fonts = properties.find(qn("w:rFonts"))
            if fonts is None:
                fonts = OxmlElement("w:rFonts")
                properties.insert(0, fonts)
            for attribute in ("ascii", "hAnsi", "eastAsia"):
                fonts.set(qn(f"w:{attribute}"), "Times New Roman")
            size = properties.find(qn("w:sz"))
            if size is None:
                size = OxmlElement("w:sz")
                properties.append(size)
            size.set(qn("w:val"), "24")
            color = properties.find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                properties.append(color)
            color.set(qn("w:val"), "000000")
            underline = properties.find(qn("w:u"))
            if underline is None:
                underline = OxmlElement("w:u")
                properties.append(underline)
            underline.set(qn("w:val"), "none")


def format_references(document: Document) -> None:
    in_references = False
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if text == "References":
            in_references = True
            paragraph.paragraph_format.keep_with_next = True
            continue
        if not in_references or not text:
            continue
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        paragraph.paragraph_format.keep_together = True


def normalize_docx(source: Path, output: Path) -> None:
    document = Document(source)
    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    normalize_styles(document)
    split_long_audit_table(document)
    split_wide_source_table(document)
    for paragraph in document.paragraphs:
        kind = style_kind(paragraph.style.name)
        paragraph.paragraph_format.line_spacing = 1.5
        if kind in {"section", "subsection", "subsubsection", "paragraph"}:
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.left_indent = Inches(0)
            paragraph.paragraph_format.right_indent = Inches(0)
            paragraph.paragraph_format.first_line_indent = Inches(0)
        for run in paragraph.runs:
            set_font(run, size=12)
            if kind in {
                "title",
                "abstract",
                "section",
                "subsection",
                "subsubsection",
                "paragraph",
            }:
                run.bold = True

    keep_captions_with_objects(document)
    resize_second_figure(document)
    format_hyperlink_runs(document)
    format_references(document)
    normalize_tables(document)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.title = (
        "From Endogenous Credit to Borrower Composition: "
        "A Reproducible Measure from Japanese Sectoral Loan Stocks"
    )
    document.core_properties.subject = ""
    document.core_properties.keywords = ""
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def m_el(local: str, **attributes: str) -> etree._Element:
    element = etree.Element(f"{{{M_NS}}}{local}")
    for key, value in attributes.items():
        element.set(f"{{{M_NS}}}{key}", value)
    return element


def math_run(text: str, *, italic: bool = False) -> etree._Element:
    run = m_el("r")
    if italic:
        properties = m_el("rPr")
        properties.append(m_el("sty", val="i"))
        run.append(properties)
    text_node = m_el("t")
    if text.startswith(" ") or text.endswith(" "):
        text_node.set(f"{{{XML_NS}}}space", "preserve")
    text_node.text = text
    run.append(text_node)
    return run


def append_math_children(container: etree._Element, node: etree._Element) -> None:
    local = etree.QName(node).localname
    children = [child for child in node if isinstance(child.tag, str)]
    if local in {"math", "mrow", "mstyle"}:
        if node.text and node.text.strip():
            container.append(math_run(" ".join(node.text.split())))
        for child in children:
            append_math_children(container, child)
            if child.tail and child.tail.strip():
                container.append(math_run(" ".join(child.tail.split())))
        return
    if local in {"mi", "mn", "mo", "mtext"}:
        text = "".join(node.itertext())
        if text:
            container.append(math_run(text, italic=(local == "mi")))
        return
    if local == "mspace":
        width = node.get("width", "")
        container.append(math_run("    " if "2em" in width else " "))
        return
    if local in {"msub", "msup", "msubsup"}:
        tag = {"msub": "sSub", "msup": "sSup", "msubsup": "sSubSup"}[local]
        construct = m_el(tag)
        construct.append(m_el(f"{tag}Pr"))
        base = m_el("e")
        append_math_children(base, children[0])
        construct.append(base)
        if local in {"msub", "msubsup"}:
            sub = m_el("sub")
            append_math_children(sub, children[1])
            construct.append(sub)
        if local in {"msup", "msubsup"}:
            sup_index = 1 if local == "msup" else 2
            sup = m_el("sup")
            append_math_children(sup, children[sup_index])
            construct.append(sup)
        container.append(construct)
        return
    if local == "mfrac":
        fraction = m_el("f")
        fraction.append(m_el("fPr"))
        numerator = m_el("num")
        append_math_children(numerator, children[0])
        denominator = m_el("den")
        append_math_children(denominator, children[1])
        fraction.extend([numerator, denominator])
        container.append(fraction)
        return
    if local == "mover":
        accent = m_el("acc")
        accent_properties = m_el("accPr")
        mark = "".join(children[1].itertext()).strip() or "^"
        accent_properties.append(m_el("chr", val=mark))
        accent.append(accent_properties)
        base = m_el("e")
        append_math_children(base, children[0])
        accent.append(base)
        container.append(accent)
        return
    if local in {"munder", "munderover"}:
        operator = "".join(children[0].itertext()).strip()
        if "∑" in operator:
            lower_construct = m_el("limLow")
            base = m_el("e")
            base.append(math_run("∑"))
            lower = m_el("lim")
            append_math_children(lower, children[1])
            lower_construct.extend([base, lower])
            if local == "munder":
                container.append(lower_construct)
                return
            upper_construct = m_el("limUpp")
            upper_base = m_el("e")
            upper_base.append(lower_construct)
            upper = m_el("lim")
            if local == "munderover":
                append_math_children(upper, children[2])
            upper_construct.extend([upper_base, upper])
            container.append(upper_construct)
        else:
            limit = m_el("limLow")
            base = m_el("e")
            append_math_children(base, children[0])
            lower = m_el("lim")
            append_math_children(lower, children[1])
            limit.extend([base, lower])
            container.append(limit)
        return

    text = "".join(node.itertext())
    if text:
        container.append(math_run(text))


def mathml_to_omml(serialized: bytes) -> etree._Element:
    source = etree.fromstring(serialized)
    equation = m_el("oMath")
    append_math_children(equation, source)
    return equation


def make_text_run(template: etree._Element, text: str) -> etree._Element:
    run = etree.Element(f"{{{W_NS}}}r")
    run_properties = template.find(f"{{{W_NS}}}rPr")
    if run_properties is not None:
        run.append(copy.deepcopy(run_properties))
    text_node = etree.SubElement(run, f"{{{W_NS}}}t")
    if text.startswith(" ") or text.endswith(" "):
        text_node.set(f"{{{XML_NS}}}space", "preserve")
    text_node.text = text
    return run


def replace_math_markers(document_xml: bytes, math_nodes: list[bytes]) -> tuple[bytes, int]:
    root = etree.fromstring(document_xml)
    replaced = 0
    for text_node in list(root.xpath("//w:t[contains(., '[[MATH')]", namespaces=NS)):
        text = text_node.text or ""
        matches = list(MATH_MARKER.finditer(text))
        if not matches:
            continue
        run = text_node.getparent()
        parent = run.getparent()
        position = parent.index(run)
        cursor = 0
        replacements: list[etree._Element] = []
        for match in matches:
            if match.start() > cursor:
                replacements.append(make_text_run(run, text[cursor:match.start()]))
            math_index = int(match.group(1))
            replacements.append(mathml_to_omml(math_nodes[math_index]))
            replaced += 1
            cursor = match.end()
        if cursor < len(text):
            replacements.append(make_text_run(run, text[cursor:]))
        parent.remove(run)
        for offset, replacement in enumerate(replacements):
            parent.insert(position + offset, replacement)

    output = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    if b"[[MATH" in output:
        raise RuntimeError("unreplaced MathML markers remain in document.xml")
    return output, replaced


def embed_external_images(
    files: dict[str, bytes],
) -> tuple[dict[str, bytes], int]:
    rels_name = "word/_rels/document.xml.rels"
    rels = etree.fromstring(files[rels_name])
    embedded = 0
    embedded_ids: set[str] = set()
    for relationship in rels:
        if not relationship.get("Type", "").endswith("/image"):
            continue
        if relationship.get("TargetMode") != "External":
            continue
        parsed = urllib.parse.urlparse(relationship.get("Target", ""))
        source = Path(urllib.parse.unquote(parsed.path))
        if not source.exists():
            raise FileNotFoundError(f"linked figure is missing: {source}")
        extension = source.suffix.lower() or ".png"
        media_name = f"word/media/ref_figure_{embedded + 1}{extension}"
        files[media_name] = source.read_bytes()
        relationship.set("Target", f"media/{Path(media_name).name}")
        relationship.attrib.pop("TargetMode", None)
        embedded_ids.add(relationship.get("Id", ""))
        embedded += 1
    files[rels_name] = etree.tostring(
        rels, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    document = etree.fromstring(files["word/document.xml"])
    for blip in document.xpath("//*[@r:link]", namespaces={"r": R_DOC_NS}):
        relationship_id = blip.get(f"{{{R_DOC_NS}}}link")
        if relationship_id not in embedded_ids:
            continue
        blip.attrib.pop(f"{{{R_DOC_NS}}}link", None)
        blip.set(f"{{{R_DOC_NS}}}embed", relationship_id)
    files["word/document.xml"] = etree.tostring(
        document, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    content_types = etree.fromstring(files["[Content_Types].xml"])
    defaults = {
        element.get("Extension", "").lower()
        for element in content_types.findall(f"{{{CT_NS}}}Default")
    }
    for extension, mime in (("png", "image/png"), ("jpg", "image/jpeg"), ("jpeg", "image/jpeg")):
        if extension in defaults:
            continue
        default = etree.Element(f"{{{CT_NS}}}Default")
        default.set("Extension", extension)
        default.set("ContentType", mime)
        content_types.append(default)
    files["[Content_Types].xml"] = etree.tostring(
        content_types, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    return files, embedded


def patch_package(source: Path, output: Path, math_nodes: list[bytes]) -> tuple[int, int]:
    with zipfile.ZipFile(source) as archive:
        files = {name: archive.read(name) for name in archive.namelist()}
    files["word/document.xml"], replaced = replace_math_markers(
        files["word/document.xml"], math_nodes
    )
    files, embedded = embed_external_images(files)
    files.pop("docProps/custom.xml", None)
    package_relationships = etree.fromstring(files["_rels/.rels"])
    for relationship in list(package_relationships):
        if relationship.get("Type", "").endswith("/custom-properties"):
            package_relationships.remove(relationship)
    files["_rels/.rels"] = etree.tostring(
        package_relationships,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )
    content_types = etree.fromstring(files["[Content_Types].xml"])
    for override in list(content_types):
        if override.get("PartName") == "/docProps/custom.xml":
            content_types.remove(override)
    files["[Content_Types].xml"] = etree.tostring(
        content_types,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name in sorted(files):
            archive.writestr(name, files[name])
    return replaced, embedded


def main() -> int:
    args = parse_args()
    root = Path(__file__).resolve().parents[1]
    soffice = locate_soffice(args.soffice)
    with tempfile.TemporaryDirectory(prefix="ref-docx-", dir="/private/tmp") as tmp:
        work = Path(tmp)
        html_source = convert_tex_to_html(root / "tex", work)
        marked_html, math_nodes = prepare_marked_html(html_source)
        converted = convert_html_to_docx(marked_html, work, soffice)
        normalized = work / "normalized.docx"
        normalize_docx(converted, normalized)
        replaced, embedded = patch_package(normalized, args.output, math_nodes)
    if replaced != len(math_nodes):
        raise RuntimeError(
            f"replaced {replaced} equations, but TeX4ht emitted {len(math_nodes)}"
        )
    print(
        f"built {args.output} with {replaced} editable equations "
        f"and {embedded} embedded figures"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
